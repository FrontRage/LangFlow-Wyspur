from docx import Document

def docx_to_string(filepath):
    # Load the .docx file
    doc = Document(filepath)
    
    # Extract all text from paragraphs
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    # Join all paragraphs into a single string separated by newlines
    return "\n".join(full_text)

import tiktoken

def count_tokens(text: str, model: str = "gpt-3.5-turbo") -> int:
    """
    Count the number of tokens in the given text for a specified model.
    
    Parameters:
        text (str): The text to be tokenized.
        model (str): The name of the model (e.g. "gpt-3.5-turbo", "gpt-4").
    
    Returns:
        int: The number of tokens.
    """
    try:
        # For newer models like GPT-4 and GPT-3.5-turbo, use tiktoken's encoding_for_model
        encoding = tiktoken.encoding_for_model(model)
    except KeyError:
        # If the model is not recognized by encoding_for_model, fall back to a default encoding
        # or specify a known encoding directly.
        encoding = tiktoken.get_encoding("cl100k_base")
    
    tokens = encoding.encode(text)
    return len(tokens)


# Example usage:
if __name__ == "__main__":
    sample_text = "This is a sample transcript text that could come from a docx file."
    
    token_count_3_5 = count_tokens(sample_text, model="gpt-3.5-turbo")
    token_count_4 = count_tokens(sample_text, model="gpt-4")
    
    print(f"Tokens with gpt-3.5-turbo: {token_count_3_5}")
    print(f"Tokens with gpt-4: {token_count_4}")

def export_summary_to_docx(summary_text, output_file_path):
    # Create a new Word document
    doc = Document()
    
    # Add a title
    doc.add_heading('Meeting Summary', level=1)
    
    # Add the summary text
    doc.add_paragraph(summary_text)
    
    # Save the document
    doc.save(output_file_path)